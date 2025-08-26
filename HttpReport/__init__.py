import os, io, json, tempfile, logging
import azure.functions as func
import pandas as pd
from docx import Document
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from pypdf import PdfReader, PdfWriter
from azure.storage.blob import BlobServiceClient, generate_blob_sas, BlobSasPermissions
from datetime import datetime, timedelta
from requests_toolbelt.multipart.decoder import MultipartDecoder

# ENV VARS
CONN_STR = os.environ.get("AZURE_STORAGE_CONNECTION_STRING", "")
CONTAINER = os.environ.get("BLOB_CONTAINER", "reports")

def _parse_multipart(req: func.HttpRequest):
    ctype = req.headers.get('Content-Type') or req.headers.get('content-type')
    if not ctype or "multipart/form-data" not in ctype:
        raise ValueError("Content-Type must be multipart/form-data")
    dec = MultipartDecoder(req.get_body(), ctype)
    files = {}
    for part in dec.parts:
        cd = part.headers.get(b'Content-Disposition', b'').decode("utf-8", errors="ignore")
        name = None; filename = None
        for token in cd.split(";"):
            token = token.strip()
            if token.startswith("name="):
                name = token.split("=",1)[1].strip('"')
            if token.startswith("filename="):
                filename = token.split("=",1)[1].strip('"')
        if name:
            files[name] = {"filename": filename, "content": part.content}
    return files

def _read_scores_from_excel(excel_bytes: bytes):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as f:
        f.write(excel_bytes); f.flush()
        df = pd.read_excel(f.name, sheet_name="Summary Dashboard")
    cols_map = {c.strip().lower(): c for c in df.columns if isinstance(c,str)}
    cat_col = cols_map.get("category")
    avg_col = cols_map.get("average score") or list(df.columns)[1]
    if not cat_col:
        raise ValueError("Could not find 'Category' column in Summary Dashboard.")
    df = df.dropna(subset=[cat_col])
    def to_num(v): 
        try: return float(v)
        except: return None
    cats, scores, overall = [], [], None
    for _, r in df.iterrows():
        name = str(r[cat_col]).strip()
        val = to_num(r.get(avg_col))
        if not name: 
            continue
        if "overall" in name.lower():
            if val is not None: overall = val
            continue
        if val is not None:
            cats.append(name)
            scores.append(val)
    return cats, scores, overall

def _build_scorecard_pdf(categories, scores, overall, out_path):
    from reportlab.lib.styles import getSampleStyleSheet
    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(out_path, pagesize=letter)
    elements = []
    elements.append(Paragraph("CJ Strategies Hospitality Consulting<br/>Boutique Hotel Operations Scorecard", styles["Title"]))
    elements.append(Spacer(1, 20))
    data = [["Category", "Average Score (1-5)"]]
    for c, s in zip(categories, scores):
        data.append([c, round(s,2)])
    data.append(["Overall Average", round(overall,2) if overall is not None else ""])
    table = Table(data, colWidths=[250, 200])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#1C2A39")),
        ("TEXTCOLOR", (0,0), (-1,0), colors.white),
        ("ALIGN", (0,0), (-1,-1), "CENTER"),
        ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
        ("BOTTOMPADDING", (0,0), (-1,0), 12),
        ("GRID", (0,0), (-1,-1), 0.5, colors.grey),
        ("BACKGROUND", (0,1), (-1,-1), colors.HexColor("#F7F6F3")),
    ]))
    elements += [table, Spacer(1,20), Paragraph("CJ Strategies Hospitality Consulting – Boutique Hotel Operations, Elevated.", styles["Normal"])]
    doc.build(elements)

def _fill_word_scores(template_bytes: bytes, score_map: dict) -> bytes:
    tmp_in = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp_in.write(template_bytes); tmp_in.close()
    doc = Document(tmp_in.name)
    if not doc.tables:
        raise ValueError("Word template has no tables. Expected Operations Review table as first table.")
    t = doc.tables[0]
    norm = {k.strip().lower(): v for k, v in score_map.items()}
    for row in t.rows[1:]:
        key = row.cells[0].text.strip().lower()
        if key in norm:
            row.cells[1].text = f"{norm[key]:.2f}"
    tmp_out = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp_out.name); tmp_out.close()
    with open(tmp_out.name, "rb") as f:
        data = f.read()
    os.unlink(tmp_in.name); os.unlink(tmp_out.name)
    return data

def _merge_pdfs(cover_pdf_bytes: bytes, scorecard_pdf_path: str, narrative_pdf_bytes: bytes|None) -> bytes:
    writer = PdfWriter()
    if cover_pdf_bytes:
        cover = PdfReader(io.BytesIO(cover_pdf_bytes))
        for p in cover.pages: writer.add_page(p)
    score = PdfReader(scorecard_pdf_path)
    for p in score.pages: writer.add_page(p)
    if narrative_pdf_bytes:
        narr = PdfReader(io.BytesIO(narrative_pdf_bytes))
        for p in narr.pages: writer.add_page(p)
    out = io.BytesIO()
    writer.write(out); out.seek(0)
    return out.getvalue()

def _upload_with_sas(name: str, data: bytes):
    if not CONN_STR:
        raise ValueError("AZURE_STORAGE_CONNECTION_STRING not configured.")
    bsc = BlobServiceClient.from_connection_string(CONN_STR)
    container = bsc.get_container_client(CONTAINER)
    try:
        container.create_container()
    except Exception:
        pass
    container.upload_blob(name, data, overwrite=True)
    acc_name = bsc.account_name
    acc_key = bsc.credential.account_key
    sas = generate_blob_sas(
        account_name=acc_name,
        container_name=CONTAINER,
        blob_name=name,
        account_key=acc_key,
        permission=BlobSasPermissions(read=True),
        expiry=datetime.utcnow() + timedelta(hours=24)
    )
    return f"https://{acc_name}.blob.core.windows.net/{CONTAINER}/{name}?{sas}"

def main(req: func.HttpRequest) -> func.HttpResponse:
    try:
        files = _parse_multipart(req)
        excel = files.get("excel")
        word_template = files.get("word_template")
        cover_pdf = files.get("cover_pdf")
        narrative_pdf = files.get("narrative_pdf")
        if not excel or not word_template:
            return func.HttpResponse("Missing 'excel' or 'word_template'", status_code=400)
        excel_bytes = excel["content"]
        word_bytes = word_template["content"]
        cover_bytes = cover_pdf["content"] if cover_pdf else None
        cats, scores, overall = _read_scores_from_excel(excel_bytes)
        score_map = dict(zip(cats, scores))
        tmp_score_path = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf").name
        _build_scorecard_pdf(cats, scores, overall, tmp_score_path)
        filled_docx_bytes = _fill_word_scores(word_bytes, score_map)
        narr_bytes = narrative_pdf["content"] if narrative_pdf else None
        final_pdf = _merge_pdfs(cover_bytes, tmp_score_path, narr_bytes)
        ts = datetime.utcnow().strftime("%Y%m%d-%H%M%S")
        pdf_name = f"CJ_Report_{ts}.pdf"
        docx_name = f"CJ_Narrative_{ts}.docx"
        pdf_url = _upload_with_sas(pdf_name, final_pdf)
        docx_url = _upload_with_sas(docx_name, filled_docx_bytes)
        body = {"pdf_url": pdf_url, "narrative_docx_url": docx_url}
        return func.HttpResponse(json.dumps(body), headers={"Content-Type": "application/json"}, status_code=200)
    except Exception as e:
        logging.exception("CJ Azure Function Error")
        return func.HttpResponse(str(e), status_code=500)
